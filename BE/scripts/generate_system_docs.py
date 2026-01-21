from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_document():
    document = Document()

    # Style configuration
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title
    title = document.add_heading('TÀI LIỆU HỆ THỐNG MEDICAL WEBSITE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- SECTION 1: OPERATION FLOWS ---
    document.add_heading('1. Sơ đồ Luồng Hoạt động (Operation Flows)', level=1)

    # 1.1 Auth
    document.add_heading('1.1. Luồng Đăng ký & Đăng nhập', level=2)
    p = document.add_paragraph()
    p.add_run('Khách hàng (Customer):').bold = True
    document.add_paragraph('1. Truy cập trang Đăng ký.', style='List Number')
    document.add_paragraph('2. Nhập thông tin (Tên đăng nhập, Mật khẩu, Email/SĐT).', style='List Number')
    document.add_paragraph('3. Hệ thống tạo tài khoản và tự động gán vai trò "Customer".', style='List Number')
    document.add_paragraph('4. Đăng nhập bằng Email/SĐT/Username.', style='List Number')

    p = document.add_paragraph()
    p.add_run('Nhân viên (Staff):').bold = True
    document.add_paragraph('1. Admin tạo tài khoản từ Admin Panel.', style='List Number')
    document.add_paragraph('2. Gán vai trò (Admin, Staff, WarehouseStaff) và Kho làm việc.', style='List Number')
    document.add_paragraph('3. Admin/WarehouseStaff đăng nhập Admin Panel.', style='List Number')
    document.add_paragraph('4. Staff bán hàng đăng nhập Website (POS).', style='List Number')

    # 1.2 Online Order
    document.add_heading('1.2. Luồng Mua hàng Online', level=2)
    document.add_paragraph('1. Khách hàng thêm sản phẩm vào Giỏ hàng.', style='List Number')
    document.add_paragraph('2. Thanh toán (Checkout) -> Chọn địa chỉ, mốc liều (COD/VNPay).', style='List Number')
    document.add_paragraph('3. Hệ thống tạo SaleInvoice (Status: Pending).', style='List Number')
    document.add_paragraph('4. Admin duyệt đơn -> Chuyển trạng thái Shipping.', style='List Number')
    document.add_paragraph('5. Giao hàng thành công -> Status: Delivered.', style='List Number')

    # 1.3 In-Store Sale
    document.add_heading('1.3. Luồng Bán hàng Tại quầy (POS)', level=2)
    document.add_paragraph('1. Nhân viên đăng nhập Website.', style='List Number')
    document.add_paragraph('2. Chọn Khách hàng (hoặc Khách lẻ).', style='List Number')
    document.add_paragraph('3. Chọn sản phẩm từ Kho chi nhánh (User.warehouseId).', style='List Number')
    document.add_paragraph('4. Thanh toán & Xuất hóa đơn.', style='List Number')
    document.add_paragraph('5. Hệ thống trừ tồn kho chi nhánh, tạo SaleInvoice (Status: Delivered).', style='List Number')

    # 1.4 Inventory
    document.add_heading('1.4. Luồng Quản lý Kho', level=2)
    document.add_paragraph('1. Nhập kho (Import): Hàng về Kho tổng (Central).', style='List Number')
    document.add_paragraph('2. Chuyển kho (Transfer): Admin tạo phiếu chuyển từ Kho tổng -> Kho chi nhánh.', style='List Number')
    document.add_paragraph('3. Kho chi nhánh nhận hàng -> Cập nhật tồn kho.', style='List Number')

    # --- SECTION 2: DATABASE SCHEMA ---
    document.add_heading('2. Cơ sở Dữ liệu & Ý nghĩa các Bảng', level=1)

    def add_table(headers, data):
        table = document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = None # Default color
        
        for row_data in data:
            row_cells = table.add_row().cells
            for i, item in enumerate(row_data):
                row_cells[i].text = item

    # 2.1 Auth
    document.add_heading('2.1. Phân quyền (Auth)', level=2)
    data_auth = [
        ['User', 'Thông tin người dùng (Khách/NV)', 'userName, passWord, email, phoneNum, roleId, warehouseId'],
        ['Role', 'Vai trò hệ thống', 'roleName (Admin, Staff...), description']
    ]
    add_table(['Bảng', 'Ý nghĩa', 'Các trường chính'], data_auth)

    # 2.2 Product
    document.add_heading('2.2. Sản phẩm (Product)', level=2)
    data_prod = [
        ['Product', 'Thông tin thuốc', 'productName, price, stockQuantity, unit, variants'],
        ['Category', 'Danh mục', 'categoryName'],
        ['Manufacturer', 'Nhà sản xuất', 'manufacturerName'],
        ['ProductBatch', 'Lô hàng (Hạn sử dụng)', 'batchNumber, expiryDate, quantity, warehouseId']
    ]
    add_table(['Bảng', 'Ý nghĩa', 'Các trường chính'], data_prod)

    # 2.3 Warehouse
    document.add_heading('2.3. Kho hàng (Warehouse)', level=2)
    data_wh = [
        ['Warehouse', 'Danh sách kho', 'warehouseName, address, warehouseType (central/branch)'],
        ['InventoryTransfer', 'Phiếu chuyển kho', 'fromWarehouseId, toWarehouseId, products, status']
    ]
    add_table(['Bảng', 'Ý nghĩa', 'Các trường chính'], data_wh)

    # 2.4 Orders
    document.add_heading('2.4. Đơn hàng (Invoice)', level=2)
    data_inv = [
        ['SaleInvoice', 'Hóa đơn bán hàng', 'userId, staffId, totalAmount, paymentMethod, statusId'],
        ['SaleInvoiceDetail', 'Chi tiết đơn hàng', 'productId, quantity, unitPrice'],
        ['OrderStatus', 'Trạng thái đơn hàng', 'statusName (Pending, Shipping...)']
    ]
    add_table(['Bảng', 'Ý nghĩa', 'Các trường chính'], data_inv)

    # 2.5 Purchases
    document.add_heading('2.5. Nhập hàng (Purchase)', level=2)
    data_pur = [
        ['PurchaseInvoice', 'Hóa đơn nhập hàng', 'staffId, supplierName, importDate, totalAmount'],
        ['PurchaseInvoiceDetail', 'Chi tiết nhập', 'productId, quantity, unitPrice, totalPrice']
    ]
    add_table(['Bảng', 'Ý nghĩa', 'Các trường chính'], data_pur)

    # 2.6 System & Others
    document.add_heading('2.6. Hệ thống & Khác', level=2)
    data_other = [
        ['Setting', 'Cấu hình hệ thống', 'storeName, phone, email, notifyLowStock...'],
        ['OTP', 'Mã xác thực (Reset Password)', 'email, otp, createdAt']
    ]
    add_table(['Bảng', 'Ý nghĩa', 'Các trường chính'], data_other)

    # Save
    file_path = 'Tai_Lieu_He_Thong_Medical_Full.docx'
    document.save(file_path)
    print(f"Successfully created {file_path}")

if __name__ == '__main__':
    create_document()
